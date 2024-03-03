from docx import Document
from docx.shared import Inches

import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

name = 'Nelson'
phone_number ='000000'
email ='hello@test.com'

document.add_picture(
    '2x2.jpg',
    width=Inches(2.0)

)

document.add_paragraph(
    name+ ' |   ' + phone_number    +'  |   '+email

)

#about me 
document.add_heading('About me')
about_me = 'Tell me about yourself'
document.add_paragraph(about_me)

#work experiences
document.add_heading('Work experience')
p = document.add_paragraph()

company ='Seven Seven'
from_date = "05/13/2015"
to_date = 'Present'

p.add_run(company+ " ").bold = True
p.add_run(from_date+'-'+to_date+'\n').italic= True
experience_details = 'Test experiences '+ company


while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company ='Seven Seven'
        from_date = "05/13/2015"
        to_date = 'Present'

        p.add_run(company+ " ").bold = True
        p.add_run(from_date+'-'+to_date+'\n').italic= True
        experience_details = 'Test experiences '+ company


        p.add_run(experience_details)
    else:
        break


document.add_heading('Skills')
skills =[];
while True:
    has_skills = input('Do you have skills? Yes or no ')
    if has_skills.lower() ==  'yes':
        skill_name = input('Enter your skill: ')
        skills.append(skill_name.capitalize())
    else:
        break

#p = document.add_paragraph()
#p.style = 'List Bullet'


for skill in skills:
   
    document.add_paragraph(skill, style='List Bullet')



document.save('cv.docx')